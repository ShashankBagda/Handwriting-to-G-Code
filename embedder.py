from docx import Document
from docx.shared import Inches

# Function to read text from a Docx file
def read_text_from_docx(file_path):
    doc = Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

# List of special characters and capital letters
names = ['_full_stop','_double_InComma','_underscore','_equals','_addition','_substraction','_brac_curl_end','_brac_curl_start','_brac_square_end','_brac_square_start','_multiplication','_hashtag','_and','_percentage','_dollar','_tilde','_brac_round_end','_brac_round_start','_greaterthan','_lessthan','_slash','_toPower','_atharet','_question','_exclamation','_colon','_comma','_apostrophe','_backquote','_semicolon','_9', '_8', '_7', '_6', '_5', '_4', '_3', '_2', '_1', '_0','z','y', 'x', 'w', 'v', 'u', 't', 's', 'r', 'q', 'p', 'o', 'n', 'm', 'l', 'k', 'j', 'i', 'h', 'g', 'f', 'e', 'd', 'c', 'b', 'a', '_Z', '_Y', '_X', '_W', '_V', '_U', '_T', '_S', '_R', '_Q', '_P', '_O', '_N', '_M', '_L', '_K', '_J', '_I', '_H', '_G', '_F', '_E', '_D', '_C', '_B', '_A']

# Mapping special characters
special_chars = {
    '.': '_full_stop',
    '"': '_double_InComma',
    '_': '_underscore',
    '=': '_equals',
    '+': '_addition',
    '-': '_substraction',
    '}': '_brac_curl_end',
    '{': '_brac_curl_start',
    ']': '_brac_square_end',
    '[': '_brac_square_start',
    '*': '_multiplication',
    '#': '_hashtag',
    '&': '_and',
    '%': '_percentage',
    '$': '_dollar',
    '~': '_tilde',
    ')': '_brac_round_end',
    '(': '_brac_round_start',
    '>': '_greaterthan',
    '<': '_lessthan',
    '/': '_slash',
    '^': '_toPower',
    '@': '_atharet',
    '?': '_question',
    '!': '_exclamation',
    ':': '_colon',
    ',': '_comma',
    "'": '_apostrophe',
    '`': '_backquote',
    ';': '_semicolon'
}

# print("Special Characters Mapping:")
# for char, name in special_chars.items():
#     print(f"{char} => {name}")



# Mapping special characters
special_chars = {name[1:]: name for name in names}

# Mapping capital characters
capital_chars = {name[1:]: name for name in names if name.startswith('_')}

input_docx = "Input.docx"
text = read_text_from_docx(input_docx)
output_doc = Document()

# Add a paragraph if there is none in the document
if not output_doc.paragraphs:
    output_doc.add_paragraph()

for char in text:
    print(f"Processing character: '{char}'")
    if char == '\n':
        output_doc.add_paragraph()
        continue
    elif char == ' ':
        file_name = f"Graphics/space.png"
        # continue
    elif char.isupper():
        file_name = f"Output/_{char}.png"
        width = Inches(0.3) 
    elif char.islower():
        file_name = f"Output/{char}.png"
        width = Inches(0.2)  
    elif char.isnumeric():
        file_name = f"Output/_{char}.png"
        width = Inches(0.4) 
    elif char in special_chars:
        name = special_chars[char]
        print(f"Found special character: '{char}', corresponding name: '{name}'")
        file_name = f"Output/{name}.png"
        width = Inches(0.35)
    else:
        # Skip characters not in special_chars or capital_chars
        continue

    # Add picture to the document
    print(f"Adding picture: {file_name}")
    output_doc.paragraphs[-1].add_run().add_picture(file_name, width=width)

output_doc.save("Output.docx")

print("Words Embedded Successfully...")
